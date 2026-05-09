"""
document_ingest_service.py — 既存資料 (PDF / DOCX / HTML / MD / TXT) のアップロード + テキスト抽出

ユーザーが過去の提案書・要件定義書などをアップロード → テキスト抽出 → knowledge artifact 化。
これにより hearing / requirements / proposal / estimate / template-builder の AI が
過去資料を「参考」として参照できるようになる。

サポート形式:
  - PDF (pypdf)
  - DOCX (python-docx)
  - HTML (BeautifulSoup)
  - Markdown (.md plain)
  - TXT (.txt plain)
"""
from __future__ import annotations

import io
import json
import re
from pathlib import Path
from typing import Optional


ALLOWED_EXT = {"pdf", "docx", "doc", "html", "htm", "md", "markdown", "txt"}
MAX_BYTES = 30 * 1024 * 1024  # 30 MB


def _detect_kind(filename: str, content_type: str | None) -> str:
    name = (filename or "").lower()
    ext = name.rsplit(".", 1)[-1] if "." in name else ""
    if ext in ALLOWED_EXT:
        return ext
    if content_type:
        if "pdf" in content_type: return "pdf"
        if "wordprocessingml" in content_type or "msword" in content_type: return "docx"
        if "html" in content_type: return "html"
        if "markdown" in content_type or "x-markdown" in content_type: return "md"
        if "plain" in content_type: return "txt"
    return ext or "txt"


# ──────────────────────────────────────────
# テキスト抽出
# ──────────────────────────────────────────
def _extract_pdf(content: bytes) -> str:
    try:
        from pypdf import PdfReader
    except ImportError:
        return "(pypdf 未インストール: テキスト抽出スキップ)"
    try:
        reader = PdfReader(io.BytesIO(content))
        pages = []
        for p in reader.pages:
            try:
                pages.append(p.extract_text() or "")
            except Exception:
                continue
        return "\n\n".join(pages).strip()
    except Exception as e:
        return f"(PDF 解析失敗: {e})"


def _extract_docx(content: bytes) -> str:
    try:
        from docx import Document
    except ImportError:
        return "(python-docx 未インストール: テキスト抽出スキップ)"
    try:
        doc = Document(io.BytesIO(content))
        parts = []
        for p in doc.paragraphs:
            t = (p.text or "").strip()
            if t:
                parts.append(t)
        # テーブル
        for table in doc.tables:
            for row in table.rows:
                row_txt = " | ".join((cell.text or "").strip() for cell in row.cells)
                if row_txt.strip():
                    parts.append(row_txt)
        return "\n".join(parts).strip()
    except Exception as e:
        return f"(DOCX 解析失敗: {e})"


def _extract_html(content: bytes) -> str:
    try:
        from bs4 import BeautifulSoup
    except ImportError:
        return content.decode("utf-8", errors="ignore")
    try:
        soup = BeautifulSoup(content, "html.parser")
        for tag in soup(["script", "style", "noscript"]):
            tag.decompose()
        text = soup.get_text(separator="\n", strip=True)
        return re.sub(r"\n{3,}", "\n\n", text)
    except Exception as e:
        return f"(HTML 解析失敗: {e})"


def _extract_text(content: bytes) -> str:
    return content.decode("utf-8", errors="ignore").strip()


def extract_text(filename: str, content: bytes, content_type: str | None) -> dict:
    """ファイルからテキストを抽出して dict で返す。"""
    if len(content) > MAX_BYTES:
        raise ValueError(f"file too large: {len(content)} bytes (max {MAX_BYTES})")

    kind = _detect_kind(filename, content_type)
    if kind == "pdf":
        text = _extract_pdf(content)
    elif kind in ("docx", "doc"):
        text = _extract_docx(content)
    elif kind in ("html", "htm"):
        text = _extract_html(content)
    elif kind in ("md", "markdown", "txt"):
        text = _extract_text(content)
    else:
        text = _extract_text(content)

    return {
        "kind": kind,
        "filename": filename,
        "size_bytes": len(content),
        "char_count": len(text),
        "text": text,
        "preview": text[:500],
    }


# ──────────────────────────────────────────
# ファイル本体の保存 (Supabase Storage / ローカルフォールバック)
# ──────────────────────────────────────────
async def store_original(account_id: int, filename: str, content: bytes, content_type: str) -> dict:
    from services import upload_service as up
    # upload_service の image アップロード経路を流用 (kind="other" で受ける)
    return await up.upload_image(
        account_id=account_id,
        kind="other",
        filename=filename,
        content=content,
        content_type=content_type or "application/octet-stream",
    )


# ──────────────────────────────────────────
# knowledge artifact として登録
# ──────────────────────────────────────────
async def ingest(
    account_id: int,
    filename: str,
    content: bytes,
    content_type: str | None,
    *,
    doc_type: str = "reference",       # proposal_reference / requirements_reference / generic / etc.
    title: str | None = None,
    tags: list[str] | None = None,
) -> dict:
    """ファイルをアップロード + テキスト抽出 + knowledge artifact 化。"""
    from services import artifact_service as art

    # テキスト抽出 (まず先に)
    extracted = extract_text(filename, content, content_type)

    # 元ファイル本体を Storage に保存 (image upload service を流用)
    try:
        stored = await store_original(account_id, filename, content, content_type or "application/octet-stream")
    except Exception as e:
        # Storage 失敗してもテキストは保存
        stored = {"url": None, "path": None, "error": str(e)}

    artifact_title = title or f"参考資料: {filename}"
    artifact_tags = list(set(["knowledge", "reference", doc_type] + (tags or [])))

    artifact = await art.create_artifact(
        type="knowledge",
        title=artifact_title,
        data={
            "doc_type": doc_type,
            "filename": filename,
            "kind": extracted["kind"],
            "size_bytes": extracted["size_bytes"],
            "char_count": extracted["char_count"],
            "preview": extracted["preview"],
            "text": extracted["text"][:50000],  # 50,000 文字まで保存
            "stored_url": stored.get("url"),
            "stored_path": stored.get("path"),
            "uploaded_at": __import__("datetime").datetime.utcnow().isoformat() + "Z",
        },
        category_tags=artifact_tags,
        created_by="user",
        actor="user",
    )
    return {
        "artifact": artifact,
        "extracted": {k: v for k, v in extracted.items() if k != "text"},  # text は重いので返さない
        "stored": stored,
    }


async def list_references(
    account_id: int | None = None,
    doc_type: str | None = None,
    limit: int = 100,
) -> list[dict]:
    """登録済み参考資料 artifact 一覧を返す。"""
    from services import artifact_service as art
    items = await art.list_artifacts(limit=limit)
    out = []
    for a in items or []:
        if a.get("type") != "knowledge":
            continue
        tags = a.get("category_tags") or []
        if "reference" not in tags:
            continue
        if doc_type and doc_type not in tags:
            continue
        data = a.get("data") or {}
        out.append({
            "id": a.get("id"),
            "title": a.get("title"),
            "tags": tags,
            "doc_type": data.get("doc_type"),
            "filename": data.get("filename"),
            "kind": data.get("kind"),
            "preview": data.get("preview", "")[:300],
            "char_count": data.get("char_count", 0),
            "stored_url": data.get("stored_url"),
            "uploaded_at": data.get("uploaded_at"),
        })
    return out


async def get_reference_text(artifact_id: str) -> str:
    """artifact からフルテキストを取得 (AI prompt に注入する用)。"""
    from services import artifact_service as art
    a = await art.get_artifact(artifact_id)
    if not a:
        return ""
    return ((a.get("data") or {}).get("text") or "")


async def build_references_context_block(
    account_id: int | None = None,
    doc_type: str | None = None,
    keywords: list[str] | None = None,
    limit: int = 3,
    max_chars_per_ref: int = 2500,
) -> str:
    """AI system prompt 用の参考資料コンテキストブロックを構築する。

    関連性の高い順に最大 `limit` 件、それぞれ `max_chars_per_ref` 文字まで含める。
    マッチが無ければ空文字を返す。
    """
    try:
        if keywords:
            refs = await find_relevant_references(
                keywords=keywords, doc_type=doc_type, limit=limit
            )
        else:
            refs = await list_references(
                account_id=account_id, doc_type=doc_type, limit=limit
            )
    except Exception:
        return ""

    if not refs:
        return ""

    lines = [
        "",
        "## 過去の参考資料 (アカウント所有・AI 参照可)",
        "以下は同一アカウントが過去に作成・受領した参考資料です。本案件のトーン・粒度・構成を揃える参考にしてください。クライアント情報と混同しないこと。",
        "",
    ]
    for r in refs:
        title = r.get("title") or r.get("filename") or "(無題)"
        kind = r.get("kind") or ""
        # フルテキストを取りに行く (preview は短すぎる)
        artifact_id = r.get("id")
        text = ""
        if artifact_id:
            try:
                text = await get_reference_text(str(artifact_id))
            except Exception:
                text = ""
        if not text:
            text = r.get("preview") or ""
        text = (text or "").strip()
        if len(text) > max_chars_per_ref:
            text = text[:max_chars_per_ref] + "…(以下略)"
        lines.append(f"### 参考資料: {title} ({kind})")
        lines.append(text or "(本文なし)")
        lines.append("")
    return "\n".join(lines)


async def find_relevant_references(
    keywords: list[str],
    doc_type: str | None = None,
    limit: int = 5,
) -> list[dict]:
    """キーワードでマッチする参考資料を検索 (簡易: タイトル + プレビューに含まれる)。"""
    refs = await list_references(doc_type=doc_type, limit=200)
    if not keywords:
        return refs[:limit]
    kws = [k.lower() for k in keywords]
    scored = []
    for r in refs:
        text = (r.get("title", "") + " " + r.get("preview", "")).lower()
        score = sum(1 for k in kws if k in text)
        if score > 0:
            scored.append((score, r))
    scored.sort(key=lambda x: -x[0])
    return [r for _, r in scored[:limit]]
